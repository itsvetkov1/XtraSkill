"""
Export service for generating downloadable artifacts in multiple formats.

Supports:
- Markdown (.md): Plain text with original formatting
- PDF (.pdf): Styled document using WeasyPrint and HTML templates
- Word (.docx): Microsoft Word document using python-docx
"""
from io import BytesIO
from pathlib import Path
from typing import Literal

import markdown
from jinja2 import Environment, FileSystemLoader
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models import Artifact

# Template directory
TEMPLATE_DIR = Path(__file__).parent.parent / "templates" / "artifacts"


def export_markdown(artifact: Artifact) -> BytesIO:
    """
    Export artifact as Markdown file.

    Args:
        artifact: The artifact to export

    Returns:
        BytesIO buffer containing UTF-8 encoded markdown
    """
    buffer = BytesIO()
    content = f"# {artifact.title}\n\n{artifact.content_markdown}"
    buffer.write(content.encode('utf-8'))
    buffer.seek(0)
    return buffer


def export_pdf(artifact: Artifact) -> BytesIO:
    """
    Export artifact as PDF using WeasyPrint.

    Converts markdown content to HTML, renders with Jinja2 template,
    and generates PDF with professional styling.

    Args:
        artifact: The artifact to export

    Returns:
        BytesIO buffer containing PDF binary data

    Raises:
        ImportError: If WeasyPrint/GTK not available (Windows without GTK)
    """
    try:
        # document_parser/__init__.py replaces xml.etree.ElementTree with defusedxml
        # in sys.modules. WeasyPrint needs the real stdlib module (it imports Element,
        # SubElement, tostring which defusedxml doesn't expose). Restore it here.
        import sys
        if 'defusedxml' in str(getattr(sys.modules.get('xml.etree.ElementTree'), '__file__', '')):
            del sys.modules['xml.etree.ElementTree']
        import xml.etree.ElementTree  # re-imports the real stdlib module
        from weasyprint import HTML
    except (ImportError, OSError) as e:
        raise ImportError(
            "PDF export requires WeasyPrint with GTK3. "
            "On Windows, install GTK3 or use Docker/Linux. "
            f"Original error: {e}"
        )

    # Convert markdown to HTML
    md = markdown.Markdown(extensions=['tables', 'fenced_code'])
    html_content = md.convert(artifact.content_markdown)

    # Render with Jinja2 template
    env = Environment(loader=FileSystemLoader(str(TEMPLATE_DIR)))
    template_name = f"{artifact.artifact_type.value}.html"

    # Fall back to user_stories template as default
    try:
        template = env.get_template(template_name)
    except Exception:
        template = env.get_template("user_stories.html")

    full_html = template.render(
        title=artifact.title,
        content=html_content,
        created_at=artifact.created_at
    )

    # Generate PDF
    buffer = BytesIO()
    HTML(string=full_html).write_pdf(buffer)
    buffer.seek(0)
    return buffer


def export_docx(artifact: Artifact) -> BytesIO:
    """
    Export artifact as Word document using python-docx.

    Parses markdown content and converts to Word formatting:
    - Headings (h1, h2, h3)
    - Bullet lists
    - Numbered lists
    - Checkbox items (displayed as text)
    - Regular paragraphs

    Args:
        artifact: The artifact to export

    Returns:
        BytesIO buffer containing .docx binary data
    """
    doc = Document()

    # Title
    heading = doc.add_heading(artifact.title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parse markdown content into sections
    lines = artifact.content_markdown.split('\n')

    for line in lines:
        line = line.rstrip()

        if not line:
            continue

        # Headings
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        # Numbered list
        elif line.startswith('1. ') or (len(line) > 2 and line[0].isdigit() and line[1] == '.'):
            text = line.split('. ', 1)[1] if '. ' in line else line
            doc.add_paragraph(text, style='List Number')
        # Checkbox items (convert to bullet)
        elif line.startswith('- [ ] ') or line.startswith('- [x] '):
            checkbox = '[x]' if '[x]' in line else '[ ]'
            text = line[6:]
            doc.add_paragraph(f"{checkbox} {text}", style='List Bullet')
        # Bold/emphasis lines (Given/When/Then) - strip asterisks
        elif line.startswith('**') and '**' in line[2:]:
            doc.add_paragraph(line.replace('**', ''))
        # Regular paragraph
        else:
            doc.add_paragraph(line)

    # Metadata footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(f"Generated by BA Assistant | {artifact.created_at.strftime('%Y-%m-%d %H:%M')}")
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.italic = True

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


ExportFormat = Literal["md", "pdf", "docx"]


def get_content_type(format: ExportFormat) -> str:
    """
    Get MIME type for export format.

    Args:
        format: Export format (md, pdf, docx)

    Returns:
        MIME type string
    """
    return {
        "md": "text/markdown",
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }[format]


def get_file_extension(format: ExportFormat) -> str:
    """
    Get file extension for export format.

    Args:
        format: Export format (md, pdf, docx)

    Returns:
        File extension string
    """
    return format if format != "docx" else "docx"
