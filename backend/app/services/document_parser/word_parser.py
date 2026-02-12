from io import BytesIO
from typing import Dict, Any
import docx
from fastapi import HTTPException

from .base import DocumentParser
from ..file_validator import validate_zip_bomb


class WordParser(DocumentParser):
    """Word parser with paragraph and table structure preservation."""

    def parse(self, file_bytes: bytes) -> Dict[str, Any]:
        """
        Extract text from Word document, preserving structure.

        Returns dict with:
            - text: Paragraphs and tables with structure headers
            - summary: First 5000 chars for AI context
            - metadata: Paragraph count, table count
        """
        try:
            doc = docx.Document(BytesIO(file_bytes))

            all_text = []

            # Extract paragraphs
            paragraphs = [
                para.text.strip()
                for para in doc.paragraphs
                if para.text.strip()
            ]

            all_text.extend(paragraphs)

            # Extract tables
            table_texts = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    cell_values = [cell.text.strip() for cell in row.cells]
                    row_text = "\t".join(cell_values)
                    table_rows.append(row_text)

                table_text = "\n".join(table_rows)
                table_texts.append(table_text)

            # Add tables section if tables exist
            if table_texts:
                all_text.append("\n[Tables]")
                all_text.extend(table_texts)

            # Join with double newlines to preserve document structure
            full_text = "\n\n".join(all_text)

            metadata = {
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
                "format": "Word"
            }

            return {
                "text": full_text,
                "summary": self.create_ai_summary(full_text),
                "metadata": metadata
            }

        except Exception as e:
            raise HTTPException(
                status_code=400,
                detail=f"Failed to parse Word file: {str(e)}"
            )

    def validate_security(self, file_bytes: bytes) -> None:
        """
        Validate Word file security.

        Raises HTTPException for malformed files or zip bombs.
        """
        # Check for zip bomb (DOCX is a zip archive)
        validate_zip_bomb(file_bytes)

        # Try opening to catch malformed DOCX
        try:
            doc = docx.Document(BytesIO(file_bytes))
            # If we can open it, it's valid
            _ = len(doc.paragraphs)
        except Exception as e:
            raise HTTPException(
                status_code=400,
                detail=f"Malformed Word file: {str(e)}"
            )
